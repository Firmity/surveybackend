"""
Facility Health Report renderer — landscape, editorial (MaintainX-inspired):
cream background, Work Sans, bold black headlines, blue + lime accents, big
numbers, chapter dividers per building, status stat-cards, a Gantt remediation
roadmap. Status colors (green/amber/red) are kept where they carry meaning.

PDF (fpdf2 + matplotlib charts):
  Cover -> Overview -> Executive Summary -> [Building divider -> category pages]*
        -> Corrective Action Plan (Gantt) -> Key Recommendations -> Back cover
DOCX path is unchanged (structured, printable).
"""
from __future__ import annotations

import base64
import io
import logging
import os
import threading
from datetime import datetime
from typing import Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from PIL import Image as PILImage

from . import charts
from .gemini import ReportContent, SectionFinding
from .scoring import grade_value
from .report_theme import (  # editable theme: colours, category labels, section text
    INK, BLUE, LIME, CREAM, CARD, LINE, MUTED, GREEN, LIME_G, AMBER, RED, SLATE,
    BRAND, REPORT_TITLE, BACK_COVER_LINE, SECTIONS,
    CATEGORY_LABELS as DOMAIN_LABELS,
)

log = logging.getLogger("render")
_FONT_DIR = os.path.join(os.path.dirname(__file__), "..", "assets", "fonts")

PhotosMap = dict[str, list[dict[str, Any]]]

# ---- landscape A4 (palette/labels/sections come from report_theme.py) ----
PAGE_W, PAGE_H = 297, 210
M = 14
CW = PAGE_W - 2 * M
WHITE = (255, 255, 255)
F = "WS"


def _dl(slug: str) -> str:
    return DOMAIN_LABELS.get(slug, slug)


# ================================================================= template injection
# The admin PDF editor stores a template (report_template.ReportTemplate). We inject it
# into this module's theme globals just before rendering, then the existing draw code
# picks it up (functions read these names as module globals at call time). A lock
# serialises renders because we mutate module state — reports are generated
# infrequently (admin action), so serialising is cheap and race-free.
_render_lock = threading.Lock()

# Palette globals that apply_template can override, in this module (RGB tuples) and in
# the charts module (hex strings). Snapshot the built-in defaults so an empty/partial
# template always resets cleanly instead of leaking the previous render's colours.
_PALETTE = ("INK", "BLUE", "LIME", "CREAM", "CARD", "LINE", "MUTED",
            "GREEN", "LIME_G", "AMBER", "RED", "SLATE")
_DEFAULT_RGB = {k: globals()[k] for k in _PALETTE}
_DEFAULT_CHART_HEX = {k: getattr(charts, k, None) for k in _PALETTE}
_DEFAULT_BRAND = {"BRAND": BRAND, "REPORT_TITLE": REPORT_TITLE, "BACK_COVER_LINE": BACK_COVER_LINE}
_DEFAULT_SECTIONS = dict(SECTIONS)
_DEFAULT_LABELS = dict(DOMAIN_LABELS)

_OVERLAYS: list[dict] = []  # active overlay elements for the current render
_SECTION_STYLES: dict = {}  # section key -> {"accent","bg","text"} hex overrides
SPACING = 1.0               # vertical spacing multiplier for responses
_TOC_ENTRIES: list = []     # (name, level, page_no) collected during the first render pass


def _ss_col(key: str, field: str, default_rgb: tuple) -> tuple:
    """Resolve a per-section colour override (accent/bg/text) or fall back to default."""
    s = _SECTION_STYLES.get(key) or {}
    v = s.get(field)
    return _hx_rgb(v) if v else default_rgb


def _sp(v: float) -> float:
    """Scale a vertical spacing value by the template's spacing multiplier."""
    return v * SPACING


def _hx_rgb(h: str) -> tuple[int, int, int]:
    try:
        s = str(h).lstrip("#")
        return (int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))
    except (ValueError, IndexError, TypeError):
        return (0, 0, 0)


def apply_template(cfg: Optional[dict]) -> None:
    """Rebind theme globals (this module + charts) from a template dict, resetting any
    field the template omits back to the built-in default. Call inside _render_lock."""
    global BRAND, REPORT_TITLE, BACK_COVER_LINE, _OVERLAYS, SPACING  # noqa: PLW0603
    g = globals()
    cfg = cfg or {}

    pal = cfg.get("palette") or {}
    for key in _PALETTE:
        hexval = pal.get(key.lower())
        if hexval:
            rgb = _hx_rgb(hexval)
            g[key] = rgb
            if _DEFAULT_CHART_HEX.get(key) is not None:
                setattr(charts, key, "#%02x%02x%02x" % rgb)
        else:  # reset to default
            g[key] = _DEFAULT_RGB[key]
            if _DEFAULT_CHART_HEX.get(key) is not None:
                setattr(charts, key, _DEFAULT_CHART_HEX[key])

    brand = cfg.get("branding") or {}
    BRAND = brand.get("brand") or _DEFAULT_BRAND["BRAND"]
    REPORT_TITLE = brand.get("report_title") or _DEFAULT_BRAND["REPORT_TITLE"]
    BACK_COVER_LINE = brand.get("back_cover_line") or _DEFAULT_BRAND["BACK_COVER_LINE"]

    secs = cfg.get("sections") or {}
    merged = dict(_DEFAULT_SECTIONS)
    for skey in ("buildings", "corrective", "key_recs"):
        s = secs.get(skey)
        if isinstance(s, dict) and s.get("title"):
            merged[skey] = (s.get("eyebrow", ""), s.get("title", ""), s.get("description", ""))
    g["SECTIONS"] = merged

    labels = dict(_DEFAULT_LABELS)
    labels.update({k: v for k, v in (cfg.get("category_labels") or {}).items() if v})
    g["DOMAIN_LABELS"] = labels

    _OVERLAYS = list(cfg.get("overlays") or [])

    _SECTION_STYLES.clear()
    for skey, sval in (cfg.get("section_styles") or {}).items():
        if isinstance(sval, dict):
            keep = {k: v for k, v in sval.items() if v and k in ("accent", "bg", "text")}
            if keep:
                _SECTION_STYLES[skey] = keep
    try:
        SPACING = min(2.0, max(0.6, float(cfg.get("spacing") or 1.0)))
    except (TypeError, ValueError):
        SPACING = 1.0


def _reset_template() -> None:
    """Restore built-in defaults after a render so module state never leaks."""
    apply_template(None)


def _overlay_img_bytes(src: str) -> Optional[bytes]:
    if not src:
        return None
    if src.startswith("data:"):
        try:
            return base64.b64decode(src.split(",", 1)[1])
        except (ValueError, IndexError):
            return None
    return None  # http(s) srcs are passed to pdf.image directly


def _draw_overlays(pdf, page: str) -> None:
    """Paint editor overlay elements whose page anchor matches (or is 'all')."""
    for ov in _OVERLAYS:
        if ov.get("page") not in (page, "all"):
            continue
        try:
            _draw_overlay(pdf, ov)
        except Exception as e:  # noqa: BLE001 - one bad element must not fail the report
            log.warning("[OVERLAY_ERR] %s: %s", ov.get("id"), e)


def _draw_overlay(pdf, ov: dict) -> None:
    op = float(ov.get("opacity", 1.0))
    if op < 1.0:
        try:
            with pdf.local_context(fill_opacity=op, stroke_opacity=op):
                _draw_overlay_body(pdf, ov)
            return
        except Exception:  # noqa: BLE001 - opacity unsupported in this fpdf2 build
            pass
    _draw_overlay_body(pdf, ov)


def _draw_overlay_body(pdf, ov: dict) -> None:
    x, y = float(ov.get("x", 0)), float(ov.get("y", 0))
    w, h = float(ov.get("w", 40)), float(ov.get("h", 12))
    kind = ov.get("type", "text")

    if kind == "text":
        pdf.set_xy(x, y)
        pdf.set_font(F, "B" if ov.get("bold") else "", float(ov.get("size", 14)))
        pdf.set_text_color(*_hx_rgb(ov.get("color", "#0b0b0b")))
        pdf.multi_cell(w, float(ov.get("size", 14)) * 0.42, str(ov.get("text", "")),
                       align={"left": "L", "center": "C", "right": "R"}.get(ov.get("align", "left"), "L"))
    elif kind == "rect":
        fill = ov.get("fill")
        stroke = ov.get("stroke")
        style = ("F" if fill else "") + ("D" if stroke else "")
        if not style:
            return
        if fill:
            pdf.set_fill_color(*_hx_rgb(fill))
        if stroke:
            pdf.set_draw_color(*_hx_rgb(stroke))
            pdf.set_line_width(float(ov.get("stroke_w", 0.6)))
        r = float(ov.get("radius", 0))
        if r > 0:
            try:
                pdf.rect(x, y, w, h, style=style, round_corners=True, corner_radius=r)
                return
            except TypeError:
                pass
        pdf.rect(x, y, w, h, style=style)
    elif kind == "line":
        pdf.set_draw_color(*_hx_rgb(ov.get("stroke", "#0b0b0b")))
        pdf.set_line_width(float(ov.get("stroke_w", 0.6)))
        pdf.line(x, y, x + w, y + h)
    elif kind == "image":
        src = ov.get("src")
        data = _overlay_img_bytes(src)
        try:
            pdf.image(io.BytesIO(data) if data else src, x=x, y=y, w=w, h=h)
        except Exception as e:  # noqa: BLE001
            log.warning("[OVERLAY_IMG_ERR] %s", e)


def _facility_lines(survey: dict) -> list[tuple[str, str]]:
    c = survey.get("contact") or {}
    name = " ".join(filter(None, [c.get("first_name"), c.get("last_name")]))
    area = survey.get("total_area")
    return [
        ("Facility", survey.get("facility_name") or "-"),
        ("Type", (survey.get("facility_type") or "-").title()),
        ("Area", f"{area} {survey.get('area_unit') or ''}".strip() if area else "-"),
        ("Address", survey.get("facility_address") or "-"),
        ("Contact", f"{name} ({c.get('email','-')})" if name else "-"),
    ]


DEPLOY_TITLES = {"estate": "Estate Management Team", "housekeeping": "Housekeeping Deployment",
                 "technical": "Technical Manpower", "security": "Security Deployment"}


def _deploy_rows(plan: dict) -> list[tuple[str, list[tuple[str, str]]]]:
    out: list[tuple[str, list[tuple[str, str]]]] = []
    if not isinstance(plan, dict):
        return out
    for gkey, rows in plan.items():
        if not isinstance(rows, dict):
            continue
        lines: list[tuple[str, str]] = []
        for role, cell in rows.items():
            if not isinstance(cell, dict):
                continue
            nums = {k: v for k, v in cell.items() if k != "remark" and isinstance(v, (int, float)) and v}
            total = sum(nums.values())
            remark = cell.get("remark") or ""
            if total == 0 and not remark:
                continue
            parts = ", ".join(f"{k}={v}" for k, v in nums.items())
            txt = f"{parts} (Total {int(total)})" if parts else f"(Total {int(total)})"
            if remark:
                txt += f" - {remark}"
            lines.append((role, txt))
        if lines:
            out.append((DEPLOY_TITLES.get(gkey, gkey.title()), lines))
    return out


def _sev_tag(sev: str) -> str:
    return {"high": "[HIGH]", "medium": "[MEDIUM]"}.get(sev, f"[{sev.upper()}]")


def _pk(area: str, domain: str) -> str:
    return f"{area}||{domain}"


# ================================================================= colors
def _sc_rgb(s: Optional[float]) -> tuple:
    if s is None:
        return SLATE
    if s >= 85:
        return GREEN
    if s >= 70:
        return LIME_G
    if s >= 50:
        return AMBER
    return RED


def _grade_rgb(rating: str) -> tuple:
    r = str(rating or "").lower()
    if "good" in r or "excellent" in r:
        return GREEN
    if "fair" in r:
        return AMBER
    return RED


def _rating_rgb(value: Any) -> tuple:
    v = str(value or "").strip().lower()
    if v in ("good", "yes"):
        return GREEN
    if v == "satisfactory":
        return AMBER
    if v in ("unsatisfactory", "no"):
        return RED
    if v in ("n/a", "na", "not applicable"):
        return SLATE
    return INK


# ================================================================= PDF core
class _PDF(FPDF):
    chrome = True
    rtitle = REPORT_TITLE
    reyebrow = BRAND

    def header(self):
        # Capture THIS page's title + chrome at header time. The footer is emitted when
        # the NEXT page starts (by which point rtitle/chrome have advanced), so reading
        # the live values would show the next section's title / a footer on full-bleed pages.
        self._page_title = self.rtitle
        self._page_chrome = self.chrome
        self.set_fill_color(*CREAM)
        self.rect(-2, -2, PAGE_W + 4, PAGE_H + 4, style="F")
        if self.chrome:
            self.set_xy(M, 7)
            self.set_font(F, "B", 9)
            self.set_text_color(*INK)
            self.cell(80, 5, self.reyebrow)
            self.set_xy(M, 7)
            self.set_font(F, "", 9)
            self.set_text_color(*MUTED)
            self.cell(CW, 5, self.rtitle, align="R")
            self.set_draw_color(*LINE)
            self.set_line_width(0.3)
            self.line(M, 14.5, PAGE_W - M, 14.5)
        _draw_overlays(self, "all")

    def footer(self):
        if getattr(self, "_page_chrome", self.chrome):
            self.set_y(-9)
            self.set_font(F, "", 8)
            self.set_text_color(*MUTED)
            self.cell(0, 5, str(getattr(self, "_page_title", "") or ""), align="L")
            self.set_y(-9)
            self.cell(0, 5, str(self.page_no()), align="R")


def _rrect(pdf, x, y, w, h, rgb, r=3.0, style="F"):
    pdf.set_fill_color(*rgb)
    pdf.set_draw_color(*rgb)
    if r and r > 0:
        try:
            pdf.rect(x, y, w, h, style=style, round_corners=True, corner_radius=r)
            return
        except TypeError:
            pass
    pdf.rect(x, y, w, h, style=style)  # square (full-bleed backgrounds)


def _txt(pdf, w, h, text, size=10, style="", rgb=INK, x=None, align="L"):
    if x is not None:
        pdf.set_x(x)
    pdf.set_font(F, style, size)
    pdf.set_text_color(*rgb)
    pdf.multi_cell(w, h, str(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align=align)


def _need(pdf, h):
    if pdf.get_y() + h > PAGE_H - pdf.b_margin:
        _content_page(pdf, pdf.rtitle)


def _tag(pdf, x, y, text, fill, fg=WHITE, size=8, h=5.2, pad=2.4, w=None) -> float:
    pdf.set_font(F, "B", size)
    if w is None:
        w = pdf.get_string_width(str(text)) + pad * 2
    _rrect(pdf, x, y, w, h, fill, r=h / 2)
    pdf.set_xy(x, y)
    pdf.set_text_color(*fg)
    pdf.cell(w, h, str(text), align="C")
    return w


def _place_img(pdf, png, x, w, y=None) -> float:
    if not png:
        return pdf.get_y()
    try:
        im = PILImage.open(io.BytesIO(png))
        h = w * im.height / im.width
    except Exception:
        h = w * 0.6
    if y is None:
        y = pdf.get_y()
    pdf.image(io.BytesIO(png), x=x, y=y, w=w)
    return y + h


def _segbar(pdf, x, y, w, h, segs):
    total = sum(c for c, _ in segs) or 1
    cx = x
    _rrect(pdf, x, y, w, h, (235, 231, 218), r=h / 2)
    for c, rgb in segs:
        if c <= 0:
            continue
        sw = w * c / total
        pdf.set_fill_color(*rgb)
        pdf.rect(cx, y, sw, h, style="F")
        cx += sw


def _statcard(pdf, x, y, w, h, label, value, accent, vsize=26):
    _rrect(pdf, x, y, w, h, CARD, r=3)
    pdf.set_draw_color(*LINE)
    try:
        pdf.rect(x, y, w, h, style="D", round_corners=True, corner_radius=3)
    except TypeError:
        pass
    _rrect(pdf, x, y, 3, h, accent, r=1)
    pdf.set_xy(x + 7, y + 5)
    pdf.set_font(F, "B", 8)
    pdf.set_text_color(*MUTED)
    pdf.cell(w - 10, 4, str(label).upper())
    pdf.set_xy(x + 6, y + h - vsize * 0.42 - 5)
    pdf.set_font(F, "B", vsize)
    pdf.set_text_color(*INK)
    pdf.cell(w - 10, vsize * 0.42, str(value))


def _h1(pdf, text, accent=BLUE):
    _need(pdf, 14)
    y = pdf.get_y()
    _rrect(pdf, M, y + 1.5, 4, 7, accent, r=1)
    pdf.set_xy(M + 8, y)
    pdf.set_font(F, "B", 16)
    pdf.set_text_color(*INK)
    pdf.cell(0, 9, str(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)


def _content_page(pdf, title):
    pdf.chrome = True
    pdf.rtitle = title
    pdf.add_page()
    pdf.set_y(20)


def _caption(pdf, text, x=None):
    pdf.set_x(x if x is not None else M)
    pdf.set_font(F, "B", 8)
    pdf.set_text_color(*MUTED)
    pdf.cell(0, 4.6, str(text).upper(), new_x=XPos.LMARGIN, new_y=YPos.NEXT)


def _rule(pdf):
    pdf.set_draw_color(*LINE)
    pdf.set_line_width(0.4)
    pdf.line(M, pdf.get_y(), PAGE_W - M, pdf.get_y())
    pdf.ln(3)


def _subhead(pdf, title, desc="", accent=BLUE, rule=True):
    _need(pdf, 24)
    if rule:
        _rule(pdf)
    _h1(pdf, title, accent)
    if desc:
        pdf.set_x(M)
        _txt(pdf, CW, 5, desc, size=9.5, rgb=MUTED)
        pdf.ln(1)


def _section_divider(pdf, eyebrow, title, description, accent=LIME, bg=INK, key=None, text=WHITE):
    pdf.chrome = False
    pdf.add_page()
    pdf.set_auto_page_break(False)
    if key:
        _toc_mark(pdf, title, level=0)
    _rrect(pdf, -2, -2, PAGE_W + 4, PAGE_H + 4, bg, r=0)
    pdf.set_xy(M, 52)
    pdf.set_font(F, "B", 12)
    pdf.set_text_color(*accent)
    pdf.cell(0, 6, str(eyebrow).upper())
    pdf.set_xy(M, 72)
    pdf.set_font(F, "B", 40)
    pdf.set_text_color(*text)
    pdf.multi_cell(255, 17, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_xy(M, 120)
    pdf.set_font(F, "", 14)
    pdf.set_text_color(210, 212, 218)
    pdf.multi_cell(215, 7, description, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if key:
        _draw_overlays(pdf, key)
    pdf.set_auto_page_break(True, margin=14)


def _fact_tiles(pdf, facts):
    facts = [(str(a.get("value")), a.get("question", "")) for a in facts if a.get("value") not in (None, "")]
    if not facts:
        return
    cols, gap, th = 4, 5, 24
    tw = (CW - gap * (cols - 1)) / cols
    i = 0
    while i < len(facts):
        row = facts[i:i + cols]
        _need(pdf, th + 4)
        ry = pdf.get_y()
        for j, (val, lab) in enumerate(row):
            x = M + j * (tw + gap)
            _rrect(pdf, x, ry, tw, th, CARD, r=3)
            pdf.set_draw_color(*LINE)
            try:
                pdf.rect(x, ry, tw, th, style="D", round_corners=True, corner_radius=3)
            except TypeError:
                pass
            _rrect(pdf, x, ry, tw, 2.4, BLUE, r=1)
            pdf.set_xy(x + 5, ry + 6)
            pdf.set_font(F, "B", 16)
            pdf.set_text_color(*INK)
            pdf.cell(tw - 8, 8, val if len(val) <= 13 else val[:12] + "…")
            pdf.set_xy(x + 5, ry + 15)
            pdf.set_font(F, "", 7.5)
            pdf.set_text_color(*MUTED)
            pdf.multi_cell(tw - 8, 3.3, str(lab)[:52], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_y(ry + th + 4)
        i += cols


# ================================================================= data
def _group_score(answers) -> Optional[int]:
    pts = [grade_value(a.get("value")) for a in answers]
    pts = [p for p in pts if p is not None]
    return round(sum(pts) / len(pts) * 100) if pts else None


def _status_counts(answers):
    g = s = p = 0
    for a in answers:
        pt = grade_value(a.get("value"))
        if pt == 1.0:
            g += 1
        elif pt == 0.5:
            s += 1
        elif pt == 0.0:
            p += 1
    return g, s, p


def _rating_counts(groups):
    good = satis = poor = na = 0
    for gr in groups:
        for a in gr.get("answers", []):
            pt = grade_value(a.get("value"))
            v = str(a.get("value") or "").strip().lower()
            if pt == 1.0:
                good += 1
            elif pt == 0.5:
                satis += 1
            elif pt == 0.0:
                poor += 1
            elif v in ("n/a", "na", "not applicable"):
                na += 1
    return {"good": good, "satis": satis, "poor": poor, "na": na}


FACILITY_FIRST = ["Site Profile"]
FACILITY_LAST = ["Staff Profile", "Client Pain Areas", "UREST Suggestion"]


def _ordered_areas(areas):
    first = [a for a in FACILITY_FIRST if a in areas]
    last = [a for a in FACILITY_LAST if a in areas]
    mids = sorted(a for a in areas if a not in first and a not in last)
    return first + mids + last


# ================================================================= sections
def _cover(pdf, survey, content, health, photos):
    pdf.chrome = False
    pdf.add_page()
    pdf.set_auto_page_break(False)  # fixed full-page layout: never overflow onto a 2nd page
    pdf.set_xy(M, 18)
    pdf.set_font(F, "B", 11)
    pdf.set_text_color(*INK)
    pdf.cell(0, 6, BRAND)
    pdf.set_xy(M, 40)
    pdf.set_font(F, "B", 46)
    pdf.set_text_color(*INK)
    pdf.cell(0, 20, "Facility Health", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_xy(M, 62)
    pdf.set_font(F, "B", 46)
    rw = pdf.get_string_width("Report  ")
    pdf.cell(rw, 20, "Report", new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf.set_text_color(*BLUE)
    pdf.cell(60, 20, str(datetime.now().year))
    pdf.set_xy(M, 92)
    pdf.set_font(F, "", 16)
    pdf.set_text_color(*MUTED)
    pdf.cell(0, 8, survey.get("facility_name") or "Facility")

    # tile mosaic (photos + blue/lime/ink blocks)
    imgs = [it["data"] for lst in photos.values() for it in lst if it.get("data")][:4]
    tiles = []
    blocks = [BLUE, LIME, INK, BLUE, LIME]
    bi = 0
    for i in range(7):
        if imgs:
            tiles.append(("img", imgs.pop(0)))
        else:
            tiles.append(("block", blocks[bi % len(blocks)]))
            bi += 1
    ty, th, gap = 120, 72, 5
    tw = (CW - gap * 6) / 7
    for i, (kind, val) in enumerate(tiles):
        x = M + i * (tw + gap)
        if kind == "img":
            try:
                pdf.image(io.BytesIO(val), x=x, y=ty, w=tw, h=th)
            except Exception:
                _rrect(pdf, x, y=ty, w=tw, h=th, rgb=BLUE, r=5)
        else:
            _rrect(pdf, x, ty, tw, th, val, r=5)

    pdf.set_xy(M, 200)
    pdf.set_font(F, "", 9)
    pdf.set_text_color(*MUTED)
    pdf.cell(0, 5, f"AI Facility Intelligence   ·   {datetime.now().strftime('%d %B %Y')}")
    _draw_overlays(pdf, "cover")  # draw cover overlays here, on the cover page
    pdf.set_auto_page_break(True, margin=14)


def _overview(pdf, survey, content, health, actions, groups):
    _content_page(pdf, "Overview")
    _toc_mark(pdf, "Overview", level=0)
    _draw_overlays(pdf, "overview")
    _h1(pdf, "At a glance")
    pdf.set_x(M)
    _txt(pdf, CW, 5, "A snapshot of the facility's overall condition, key metrics, and how each surveyed "
                     "category scored. Detailed findings follow in the sections after this.", size=9.5, rgb=MUTED)
    pdf.ln(1)
    score = (health or {}).get("overall")
    grade = (health or {}).get("grade") or content.overall_rating

    gimg = charts.gauge(score, grade)
    gy = pdf.get_y() + 2
    _place_img(pdf, gimg, x=M + 4, w=62, y=gy)
    pdf.set_xy(M + 4, gy + 62)
    pdf.set_font(F, "B", 8)
    pdf.set_text_color(*MUTED)
    pdf.cell(62, 4, "OVERALL HEALTH SCORE", align="C")

    # rating badge + facility card to the right of the gauge
    bx = M + 82
    _rrect(pdf, bx, gy + 2, CW - 82, 26, _sc_rgb(score) if score is not None else _grade_rgb(grade), r=4)
    pdf.set_xy(bx + 8, gy + 6)
    pdf.set_font(F, "B", 9)
    pdf.set_text_color(*WHITE)
    pdf.cell(0, 5, "OVERALL RATING")
    pdf.set_xy(bx + 8, gy + 12)
    pdf.set_font(F, "B", 20)
    pdf.cell(0, 10, str(grade).upper())

    # stat cards row
    high = len([a for a in (actions or []) if a.get("severity") == "high"])
    med = len([a for a in (actions or []) if a.get("severity") == "medium"])
    doms = len((health or {}).get("domains", []))
    counts = _rating_counts(groups)
    cy = gy + 34
    cw = (CW - 82 - 3 * 6) / 4
    cards = [("Categories", doms, BLUE), ("Items graded", counts["good"] + counts["satis"] + counts["poor"], INK),
             ("High-priority", high, RED), ("Medium-priority", med, AMBER)]
    for i, (lab, val, acc) in enumerate(cards):
        _statcard(pdf, bx + i * (cw + 6), cy, cw, 26, lab, val, acc, vsize=22)

    pdf.set_y(max(gy + 70, cy + 32))
    # category bars
    rows = [(_dl(d["domain"]), d["score"], d.get("graded", 0)) for d in (health or {}).get("domains", [])]
    if rows:
        _h1(pdf, "Health score by category")
        pdf.set_x(M)
        _txt(pdf, CW, 5, "Each category scored 0-100 from the surveyor's ratings. Longer green bars are "
                         "healthier; short red bars need attention.", size=9.5, rgb=MUTED)
        pdf.ln(1)
        _need(pdf, 55)
        pdf.set_y(_place_img(pdf, charts.category_bars(rows), x=M, w=min(CW, 200)) + 3)


def _exec_summary(pdf, content, groups):
    _content_page(pdf, "Executive Summary")
    _toc_mark(pdf, "Executive Summary", level=0)
    _draw_overlays(pdf, "exec_summary")
    top = 22
    _rrect(pdf, M, top, 95, 150, INK, r=6)
    pdf.set_xy(M + 8, top + 12)
    pdf.set_font(F, "B", 28)
    pdf.set_text_color(*WHITE)
    pdf.multi_cell(80, 12, "Executive Summary", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    # rating-distribution donut on a light card (readable on the dark panel)
    counts = _rating_counts(groups)
    dimg = charts.rating_donut(counts["good"], counts["satis"], counts["poor"], counts["na"])
    card_y = top + 52
    _rrect(pdf, M + 6, card_y, 83, 90, CREAM, r=5)
    pdf.set_xy(M + 12, card_y + 5)
    pdf.set_font(F, "B", 8)
    pdf.set_text_color(*MUTED)
    pdf.cell(0, 4, "RATING DISTRIBUTION")
    if dimg:
        try:
            pdf.image(io.BytesIO(dimg), x=M + 11, y=card_y + 11, w=73)
        except Exception:
            pass
    # right: summary text
    rx = M + 106
    pdf.set_xy(rx, top + 4)
    _caption(pdf, "Summary of findings", x=rx)
    pdf.set_xy(rx, pdf.get_y() + 1)
    pdf.set_font(F, "", 15)
    pdf.set_text_color(*INK)
    pdf.multi_cell(CW - 106, 8, content.executive_summary, new_x=XPos.LMARGIN, new_y=YPos.NEXT)


def _building_divider(pdf, index, name, score, ss_key="building"):
    pdf.chrome = False
    pdf.add_page()
    pdf.set_auto_page_break(False)
    _toc_mark(pdf, name, level=1)
    _rrect(pdf, -2, -2, PAGE_W + 4, PAGE_H + 4, _ss_col(ss_key, "bg", INK), r=0)
    pdf.set_xy(M, 40)
    pdf.set_font(F, "B", 12)
    pdf.set_text_color(*_ss_col(ss_key, "accent", LIME))
    pdf.cell(0, 6, f"BUILDING {index:02d}")
    pdf.set_xy(M, 70)
    pdf.set_font(F, "B", 40)
    pdf.set_text_color(*_ss_col(ss_key, "text", WHITE))
    pdf.multi_cell(180, 18, name, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if score is not None:
        pdf.set_xy(PAGE_W - 110, 90)
        pdf.set_font(F, "B", 90)
        pdf.set_text_color(*_sc_rgb(score))
        pdf.cell(96, 40, str(score), align="R")
        pdf.set_xy(PAGE_W - 110, 138)
        pdf.set_font(F, "", 13)
        pdf.set_text_color(200, 200, 200)
        pdf.cell(96, 6, "health score / 100", align="R")
    _draw_overlays(pdf, ss_key)
    pdf.set_auto_page_break(True, margin=14)


CHIP_W = 26  # fixed chip width so status pills line up evenly

# domain -> simple line-icon glyph key (drawn with fpdf primitives, no assets needed)
_ICON_FOR = {
    "general": "info", "security": "shield", "fire_safety": "flame", "hvac": "fan",
    "electrical": "bolt", "plumbing": "drop", "civil": "building", "horticulture": "leaf",
    "housekeeping": "spark", "green_building": "leaf", "green_building_key": "leaf",
    "sop_registers": "doc", "inventory": "box", "maintenance_manager": "wrench",
    "general_maintenance": "wrench", "technical_key": "wrench",
    "client_pain_areas": "info", "urest_suggestion": "info",
}


def _cat_icon(pdf, x, y, s, domain, color=WHITE) -> None:
    """Draw a minimal line glyph for `domain` inside an s×s box at (x, y)."""
    pdf.set_draw_color(*color)
    pdf.set_line_width(max(0.45, s * 0.07))
    kind = _ICON_FOR.get(domain, "ring")

    def L(a, b, c, d):
        pdf.line(x + a * s, y + b * s, x + c * s, y + d * s)

    def E(a, b, w, h):
        pdf.ellipse(x + a * s, y + b * s, w * s, h * s, "D")

    segs = {
        "shield": [(0.5, 0.08, 0.86, 0.22), (0.86, 0.22, 0.83, 0.56), (0.83, 0.56, 0.5, 0.92),
                   (0.5, 0.92, 0.17, 0.56), (0.17, 0.56, 0.14, 0.22), (0.14, 0.22, 0.5, 0.08)],
        "flame": [(0.5, 0.06, 0.8, 0.55), (0.8, 0.55, 0.5, 0.94), (0.5, 0.94, 0.2, 0.55), (0.2, 0.55, 0.5, 0.06)],
        "bolt": [(0.56, 0.06, 0.34, 0.5), (0.34, 0.5, 0.52, 0.5), (0.52, 0.5, 0.44, 0.94),
                 (0.44, 0.94, 0.7, 0.42), (0.7, 0.42, 0.5, 0.42), (0.5, 0.42, 0.56, 0.06)],
        "spark": [(0.5, 0.1, 0.5, 0.9), (0.1, 0.5, 0.9, 0.5), (0.24, 0.24, 0.76, 0.76), (0.76, 0.24, 0.24, 0.76)],
    }
    if kind in segs:
        for a, b, c, d in segs[kind]:
            L(a, b, c, d)
    elif kind == "fan":
        E(0.4, 0.4, 0.2, 0.2); L(0.5, 0.4, 0.5, 0.1); L(0.5, 0.6, 0.82, 0.78); L(0.5, 0.6, 0.18, 0.78)
    elif kind == "drop":
        L(0.5, 0.06, 0.78, 0.5); L(0.5, 0.06, 0.22, 0.5); E(0.22, 0.4, 0.56, 0.56)
    elif kind == "building":
        pdf.rect(x + 0.22 * s, y + 0.12 * s, 0.56 * s, 0.78 * s, "D")
        L(0.4, 0.12, 0.4, 0.9); L(0.22, 0.4, 0.78, 0.4); L(0.22, 0.66, 0.78, 0.66)
    elif kind == "leaf":
        E(0.18, 0.18, 0.64, 0.64); L(0.3, 0.7, 0.7, 0.3)
    elif kind == "doc":
        pdf.rect(x + 0.26 * s, y + 0.12 * s, 0.48 * s, 0.76 * s, "D")
        L(0.34, 0.34, 0.66, 0.34); L(0.34, 0.5, 0.66, 0.5); L(0.34, 0.66, 0.58, 0.66)
    elif kind == "box":
        pdf.rect(x + 0.18 * s, y + 0.34 * s, 0.64 * s, 0.5 * s, "D")
        L(0.18, 0.34, 0.5, 0.16); L(0.82, 0.34, 0.5, 0.16); L(0.5, 0.16, 0.5, 0.34)
    elif kind == "wrench":
        E(0.14, 0.14, 0.34, 0.34); L(0.4, 0.4, 0.85, 0.85)
    elif kind == "info":
        E(0.14, 0.14, 0.72, 0.72); L(0.5, 0.44, 0.5, 0.72)
    else:
        E(0.2, 0.2, 0.6, 0.6)


def _finding_card(pdf, a) -> None:
    """One flagged response as a padded white card: severity stripe + chip + question."""
    val = str(a.get("value"))
    q = a.get("question", "") or ""
    rgb = _rating_rgb(val)
    text_x = M + 6 + CHIP_W + 4
    text_w = CW - (6 + CHIP_W + 4) - 6
    pdf.set_font(F, "", 9.5)
    nlines = max(1, int(pdf.get_string_width(q) // max(text_w - 1, 1)) + 1)
    card_h = max(nlines * 4.6 + _sp(5), 9.5)
    _need(pdf, card_h + 2)
    y = pdf.get_y()
    _rrect(pdf, M, y, CW, card_h, CARD, r=2.5)
    pdf.set_draw_color(*LINE)
    pdf.set_line_width(0.3)
    try:
        pdf.rect(M, y, CW, card_h, style="D", round_corners=True, corner_radius=2.5)
    except TypeError:
        pass
    _rrect(pdf, M, y + 1, 2.2, card_h - 2, rgb, r=1)
    _tag(pdf, M + 6, y + _sp(2.5), val, rgb, size=8, h=5.2, w=CHIP_W)
    pdf.set_xy(text_x, y + _sp(2.4))
    pdf.set_font(F, "", 9.5)
    pdf.set_text_color(*INK)
    pdf.multi_cell(text_w, 4.6, q, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_y(y + card_h + _sp(2.4))


def _building_scorecard(pdf, score, g, s, p) -> None:
    """Compact scorecard row: score + Good/Satisfactory/Unsatisfactory/Total counts."""
    total = g + s + p
    acc = _sc_rgb(score) if score is not None else BLUE
    y = pdf.get_y()
    h = 20
    _rrect(pdf, M, y, CW, h, CARD, r=4)
    pdf.set_draw_color(*LINE)
    try:
        pdf.rect(M, y, CW, h, style="D", round_corners=True, corner_radius=4)
    except TypeError:
        pass
    _rrect(pdf, M, y + 2, 3, h - 4, acc, r=1)
    pdf.set_xy(M + 8, y + 3.5)
    pdf.set_font(F, "B", 22)
    pdf.set_text_color(*acc)
    pdf.cell(26, 12, str(score if score is not None else "—"))
    pdf.set_xy(M + 8, y + h - 5.5)
    pdf.set_font(F, "", 7)
    pdf.set_text_color(*MUTED)
    pdf.cell(26, 4, "SCORE / 100")
    pdf.set_draw_color(*LINE)
    pdf.set_line_width(0.3)
    pdf.line(M + 42, y + 4, M + 42, y + h - 4)
    px = M + 52
    for lab, val, col in (("Good", g, GREEN), ("Satisfactory", s, AMBER),
                          ("Unsatisfactory", p, RED), ("Total items", total, INK)):
        pdf.set_xy(px, y + 3.5)
        pdf.set_font(F, "B", 16)
        pdf.set_text_color(*col)
        pdf.cell(46, 8, str(val))
        pdf.set_xy(px, y + h - 5.5)
        pdf.set_font(F, "", 7.5)
        pdf.set_text_color(*MUTED)
        pdf.cell(46, 4, lab.upper())
        px += 52
    pdf.set_y(y + h + _sp(5))


def _category_block(pdf, area, domain, answers, sec, photos_for):
    _need(pdf, 42)
    score = _group_score(answers)
    facts = [a for a in answers if grade_value(a.get("value")) is None and (a.get("value") or a.get("remark"))]
    graded_ans = [a for a in answers if grade_value(a.get("value")) is not None]
    g, s, p = _status_counts(answers)
    acc = _sc_rgb(score) if score is not None else BLUE

    # header band: icon badge + category name + score chip
    y = pdf.get_y()
    band_h = 15
    _rrect(pdf, M, y, CW, band_h, CARD, r=3)
    pdf.set_draw_color(*LINE)
    pdf.set_line_width(0.3)
    try:
        pdf.rect(M, y, CW, band_h, style="D", round_corners=True, corner_radius=3)
    except TypeError:
        pass
    bs = 11
    by = y + (band_h - bs) / 2
    _rrect(pdf, M + 3.5, by, bs, bs, acc, r=2.5)
    _cat_icon(pdf, M + 3.5 + bs * 0.18, by + bs * 0.18, bs * 0.64, domain, WHITE)
    pdf.set_xy(M + 3.5 + bs + 5, y + 3.6)
    pdf.set_font(F, "B", 14)
    pdf.set_text_color(*INK)
    pdf.cell(CW - 60, 8, _dl(domain))
    if score is not None:
        _tag(pdf, M + CW - 30, y + (band_h - 6) / 2, f"{score}/100", acc, size=9, h=6)
    pdf.set_y(y + band_h + _sp(4))

    if facts:
        _caption(pdf, "Key facts")
        _fact_tiles(pdf, facts)
        pdf.ln(_sp(2))

    if graded_ans:
        _caption(pdf, "Status summary")
        _segbar(pdf, M, pdf.get_y(), CW, 5, [(g, GREEN), (s, AMBER), (p, RED)])
        pdf.set_y(pdf.get_y() + 7)
        pdf.set_font(F, "", 9)
        pdf.set_text_color(*MUTED)
        pdf.cell(0, 5, f"{g} Good      {s} Satisfactory      {p} Unsatisfactory",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(_sp(2))
        flagged = [a for a in graded_ans if grade_value(a.get("value")) in (0.0, 0.5)]
        if flagged:
            pdf.set_x(M)
            pdf.set_font(F, "B", 9.5)
            pdf.set_text_color(*INK)
            pdf.cell(0, 5.5, "Needs attention", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(_sp(1.5))
            for a in flagged:
                _finding_card(pdf, a)
        if g:
            pdf.ln(_sp(1))
            pdf.set_x(M)
            _txt(pdf, CW, 4.6, f"+ {g} item(s) rated Good — no action required.", size=8.5, style="I", rgb=MUTED)

    if sec:
        _findings(pdf, "Risks", sec.risks, RED)
        _findings(pdf, "Recommendations", sec.recommendations, BLUE)
    _photos_grid(pdf, photos_for)
    pdf.ln(_sp(4))
    _rule(pdf)


def _findings(pdf, title, items, accent):
    if not items:
        return
    _need(pdf, 9)
    pdf.set_x(M)
    pdf.set_font(F, "B", 9.5)
    pdf.set_text_color(*accent)
    pdf.cell(0, 5.5, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font(F, "", 9)
    pdf.set_text_color(*INK)
    for it in items:
        _need(pdf, 6)
        yy = pdf.get_y()
        pdf.set_fill_color(*accent)
        pdf.rect(M + 1, yy + 1.7, 1.6, 1.6, style="F")
        pdf.set_xy(M + 5, yy)
        pdf.multi_cell(CW - 6, 5, it, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(1)


def _photos_grid(pdf, items):
    items = [it for it in items if it.get("data")]
    if not items:
        return
    gap = 6
    cw = (CW - 2 * gap) / 3

    def _row_h(row):
        hs = []
        for it in row:
            try:
                im = PILImage.open(io.BytesIO(it["data"]))
                hs.append(min(cw * im.height / im.width, 48))
            except Exception:
                hs.append(cw * 0.6)
        return hs, max(hs) + 7

    first_hs, first_rh = _row_h(items[:3])
    _need(pdf, 8 + first_rh)
    pdf.set_x(M)
    pdf.set_font(F, "B", 9.5)
    pdf.set_text_color(*INK)
    pdf.cell(0, 5.5, "Photos", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    i = 0
    while i < len(items):
        row = items[i:i + 3]
        hs, rowh = (first_hs, first_rh) if i == 0 else _row_h(row)
        _need(pdf, rowh)
        rowy = pdf.get_y()
        for j, it in enumerate(row):
            x = M + j * (cw + gap)
            try:
                pdf.image(io.BytesIO(it["data"]), x=x, y=rowy, w=cw, h=hs[j])
            except Exception as e:
                log.warning("[RENDER_IMG_ERR] %s", e)
            pdf.set_xy(x, rowy + hs[j] + 0.5)
            pdf.set_font(F, "", 7)
            pdf.set_text_color(*MUTED)
            pdf.multi_cell(cw, 3.2, str((it.get("question") or "Photo"))[:70], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_y(rowy + rowh)
        i += 3


def _gantt_rows(actions):
    by: dict[tuple, dict] = {}
    for ac in actions or []:
        k = (ac.get("area", "-"), ac.get("domain", ""))
        d = by.setdefault(k, {"high": 0, "med": 0})
        d["high" if ac.get("severity") == "high" else "med"] += 1
    rows = []
    for (area, dom), d in by.items():
        label = f"{_dl(dom)} · {area}"[:34]
        if d["high"]:
            start, pri, n = 0, "high", d["high"] + d["med"]
        else:
            start, pri, n = 2, "medium", d["med"]
        end = min(start + max(2, (n + 1) // 2 + 1), 7)
        rows.append((label, start, end, pri, d["high"]))
    rows.sort(key=lambda r: (0 if r[3] == "high" else 1, -r[4]))
    return [(r[0], r[1], r[2], r[3]) for r in rows]


def _corrective(pdf, actions):
    if not actions:
        return
    high = [a for a in actions if a.get("severity") == "high"]
    med = [a for a in actions if a.get("severity") == "medium"]
    _content_page(pdf, "Corrective Action Plan")
    _h1(pdf, "Priority overview", accent=RED)
    pdf.set_x(M)
    _txt(pdf, CW, 5.6,
         "Every deficiency found during the survey is captured below and prioritised by urgency. HIGH-priority "
         "items are safety- or resident-experience critical and should be closed immediately; MEDIUM-priority "
         "items are planned improvements. Below: a priority breakdown, a remediation timeline, and the full "
         "action list with the specific finding and recommended fix for each.", size=10, rgb=INK)
    pdf.ln(2)
    cy = pdf.get_y()
    _statcard(pdf, M, cy, 52, 24, "Total actions", len(actions), BLUE, vsize=22)
    _statcard(pdf, M + 57, cy, 52, 24, "High priority", len(high), RED, vsize=22)
    _statcard(pdf, M + 114, cy, 52, 24, "Medium priority", len(med), AMBER, vsize=22)
    pimg = charts.action_priority(len(high), len(med))
    if pimg:
        _place_img(pdf, pimg, x=M + 174, w=CW - 174, y=cy)
    pdf.set_y(cy + 28)

    # ---- Remediation roadmap ----
    _subhead(pdf, "Remediation roadmap",
             "A recommended schedule for closing each group of issues over the coming weeks.", accent=BLUE)
    _caption(pdf, "Timeline by category  ·  weeks from now")
    gimg = charts.gantt(_gantt_rows(actions))
    _need(pdf, 60)
    pdf.set_y(_place_img(pdf, gimg, x=M, w=min(CW, 205)) + 3)
    pdf.set_x(M)
    _txt(pdf, CW, 4.6, "Red = immediate (high priority, weeks 0-2).   Amber = planned (medium priority).   "
                       "Blue dashed line = recommended re-inspection.", size=8.5, style="I", rgb=MUTED)
    pdf.ln(2)

    # ---- Action detail ----
    _subhead(pdf, "Action detail",
             "Every corrective action with its finding and recommended fix, highest priority first.", accent=INK)
    for i, ac in enumerate(actions, 1):
        _need(pdf, 15)
        yy = pdf.get_y()
        sev = ac.get("severity", "medium")
        _tag(pdf, M, yy, sev.upper(), RED if sev == "high" else AMBER, size=8, h=5.4, w=24)
        pdf.set_xy(M + 28, yy - 0.4)
        pdf.set_font(F, "B", 10)
        pdf.set_text_color(*INK)
        pdf.multi_cell(CW - 28, 5.4, f"{i}.  {_dl(ac.get('domain',''))}  ·  {ac.get('area','')}",
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_x(M + 28)
        _txt(pdf, CW - 28, 4.8, f"Finding: {ac.get('finding')} — {ac.get('question')}", size=9, rgb=INK)
        pdf.set_x(M + 28)
        _txt(pdf, CW - 28, 4.8, f"Action: {ac.get('action')}", size=9, style="I", rgb=MUTED)
        pdf.ln(2.5)


def _key_recs(pdf, content):
    if not content.key_recommendations:
        return
    _content_page(pdf, "Key Recommendations")
    _h1(pdf, "Key Recommendations", accent=LIME_G)
    pdf.set_x(M)
    _txt(pdf, CW, 5, "The highest-impact actions to raise this facility's health score, in priority order. "
                     "Addressing these first will resolve the most critical risks found during the survey.",
         size=9.5, rgb=MUTED)
    pdf.ln(2)
    for i, rec in enumerate(content.key_recommendations, 1):
        _need(pdf, 12)
        yy = pdf.get_y()
        _rrect(pdf, M, yy, 9, 9, BLUE, r=4.5)
        pdf.set_xy(M, yy)
        pdf.set_font(F, "B", 11)
        pdf.set_text_color(*WHITE)
        pdf.cell(9, 9, str(i), align="C")
        pdf.set_xy(M + 13, yy)
        _txt(pdf, CW - 15, 5.6, rec, size=11, rgb=INK)
        pdf.ln(2)


def _back_cover(pdf, survey):
    pdf.chrome = False
    pdf.add_page()
    pdf.set_auto_page_break(False)
    _rrect(pdf, -2, -2, PAGE_W + 4, PAGE_H + 4, _ss_col("back", "bg", BLUE), r=0)
    pdf.set_xy(M, 80)
    pdf.set_font(F, "B", 34)
    pdf.set_text_color(*WHITE)
    pdf.multi_cell(200, 16, "Powered by Firmity AI", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_xy(M, 118)
    pdf.set_font(F, "", 13)
    pdf.set_text_color(*LIME)
    pdf.cell(0, 7, BACK_COVER_LINE)
    pdf.set_xy(M, 190)
    pdf.set_font(F, "", 10)
    pdf.set_text_color(230, 235, 255)
    pdf.cell(0, 5, f"{BRAND}  ·  AI Facility Intelligence")
    _draw_overlays(pdf, "back")  # draw back-cover overlays here, on the back page
    pdf.set_auto_page_break(True, margin=14)


def _toc_mark(pdf, name, level: int = 0) -> None:
    """Record a section's page number for the contents page + add a PDF bookmark."""
    _TOC_ENTRIES.append((str(name), level, pdf.page_no()))
    try:
        pdf.start_section(str(name), level=level)  # PDF outline/bookmarks (best-effort)
    except Exception:  # noqa: BLE001
        pass


def _contents_page(pdf, entries) -> None:
    """Draw a Contents page from collected (name, page) entries."""
    _content_page(pdf, "Contents")
    pdf.set_xy(M, 22)
    pdf.set_font(F, "B", 12)
    pdf.set_text_color(*BLUE)
    pdf.cell(0, 6, "IN THIS REPORT", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_x(M)
    pdf.set_font(F, "B", 30)
    pdf.set_text_color(*INK)
    pdf.cell(0, 16, "Contents", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(6)
    for name, page in entries:
        y = pdf.get_y()
        pdf.set_xy(M, y)
        pdf.set_font(F, "B", 13)
        pdf.set_text_color(*INK)
        pdf.cell(220, 9, name)
        pdf.set_xy(M, y)
        pdf.set_font(F, "", 12)
        pdf.set_text_color(*MUTED)
        pdf.cell(CW, 9, str(page), align="R")
        pdf.set_draw_color(*LINE)
        pdf.set_line_width(0.2)
        pdf.line(M, y + 11, PAGE_W - M, y + 11)
        pdf.set_y(y + 11 + _sp(3.5))


def render_pdf(content: ReportContent, survey: dict,
               photos: Optional[PhotosMap] = None, view: str = "domain",
               health: Optional[dict] = None, actions: Optional[list] = None,
               groups: Optional[list] = None, template: Optional[dict] = None) -> bytes:
    # Serialise + inject the admin template, then always restore defaults so one
    # render's theme never leaks into the next (module-global theme state).
    with _render_lock:
        try:
            apply_template(template)
            # Pass 1: render once to collect section page numbers (no contents page).
            _TOC_ENTRIES.clear()
            _render_pdf_body(content, survey, photos or {}, view, health, actions, groups or [], toc=None)
            # +1 because pass 2 inserts the contents page right after the cover.
            entries = [(n, pg + 1) for (n, lvl, pg) in _TOC_ENTRIES if lvl == 0]
            # Pass 2: render for real with the contents page.
            _TOC_ENTRIES.clear()
            return _render_pdf_body(content, survey, photos or {}, view, health, actions, groups or [], toc=entries)
        finally:
            _reset_template()


def _render_pdf_body(content, survey, photos, view, health, actions, groups, toc=None) -> bytes:
    pdf = _PDF(orientation="L", format="A4")
    pdf.set_auto_page_break(True, margin=14)
    pdf.set_margins(M, 20, M)
    for st, fn in (("", "WorkSans-Regular.ttf"), ("B", "WorkSans-Bold.ttf"), ("I", "WorkSans-Italic.ttf")):
        p = os.path.join(_FONT_DIR, fn)
        if os.path.exists(p):
            pdf.add_font(F, st, p)

    grp = {(g["area"], g["domain"]): g.get("answers", []) for g in groups}
    secidx = {(s.area, s.domain): s for s in content.sections}

    _cover(pdf, survey, content, health, photos)
    if toc:  # pass 2 only: real contents page right after the cover
        _contents_page(pdf, toc)
    _overview(pdf, survey, content, health, actions, groups)
    _exec_summary(pdf, content, groups)

    areas = _ordered_areas({a for (a, _d) in grp.keys()})
    if areas:
        _section_divider(pdf, *SECTIONS["buildings"], accent=_ss_col("buildings", "accent", LIME),
                         bg=_ss_col("buildings", "bg", INK), key="buildings",
                         text=_ss_col("buildings", "text", WHITE))
    idx = 0
    for area in areas:
        doms = [d for (a, d) in grp.keys() if a == area]
        if area != "Site Profile":
            doms = [d for d in doms if d != "general"]
        if not doms:
            continue
        idx += 1
        is_overview = area == "Site Profile"
        name = "Facility Overview" if is_overview else area
        ss_key = "facility_overview" if is_overview else "building"
        bscore = _group_score([a for d in doms for a in grp.get((area, d), [])])
        _building_divider(pdf, idx, name, bscore, ss_key)
        _content_page(pdf, name)
        bg_g = bg_s = bg_p = 0
        for d in doms:
            gg, ss, pp = _status_counts(grp.get((area, d), []))
            bg_g += gg; bg_s += ss; bg_p += pp
        if (bg_g + bg_s + bg_p) > 0:
            _building_scorecard(pdf, bscore, bg_g, bg_s, bg_p)
        for d in doms:
            _category_block(pdf, area, d, grp.get((area, d), []), secidx.get((area, d)),
                            photos.get(_pk(area, d), []))

    if actions:
        _section_divider(pdf, *SECTIONS["corrective"], accent=_ss_col("corrective", "accent", LIME),
                         bg=_ss_col("corrective", "bg", INK), key="corrective",
                         text=_ss_col("corrective", "text", WHITE))
    _corrective(pdf, actions)

    if content.key_recommendations:
        _section_divider(pdf, *SECTIONS["key_recs"], accent=_ss_col("key_recs", "accent", LIME),
                         bg=_ss_col("key_recs", "bg", BLUE), key="key_recs",
                         text=_ss_col("key_recs", "text", WHITE))
    _key_recs(pdf, content)
    _back_cover(pdf, survey)
    return bytes(pdf.output())


# ---------------------------------------------------------------- DOCX
# ---- DOCX styling primitives (mirror the PDF's editorial theme) ----
_ALIGN = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
          "right": WD_ALIGN_PARAGRAPH.RIGHT}


def _rgb(t) -> RGBColor:
    return RGBColor(int(t[0]), int(t[1]), int(t[2]))


def _hexf(rgb) -> str:
    return "%02X%02X%02X" % tuple(int(c) for c in rgb)


def _shade(cell, rgb) -> None:
    """Solid background fill on a table cell (python-docx has no high-level API)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), _hexf(rgb))
    cell._tc.get_or_add_tcPr().append(shd)


def _run_shade(run, rgb) -> None:
    """Background fill behind a run — used for coloured 'chip' text (severity/status)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), _hexf(rgb))
    run._r.get_or_add_rPr().append(shd)


def _no_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    table._tbl.tblPr.append(borders)


def _cell(cell, text, *, bold=False, color=INK, size=10.0, align="left", fill=None) -> None:
    if fill is not None:
        _shade(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = _ALIGN[align]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = _rgb(color)


def _para(doc, text="", *, bold=False, italic=False, color=INK, size=10.5, align="left",
          space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = _ALIGN[align]
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(str(text))
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = _rgb(color)
    return p


def _heading(doc, text, *, size=16.0, color=INK, accent=None, space_before=10, space_after=4):
    """Editorial heading: optional coloured accent bar + bold title."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if accent is not None:
        bar = p.add_run("▍ ")  # left bar glyph as a coloured accent
        bar.font.size = Pt(size)
        bar.font.color.rgb = _rgb(accent)
    r = p.add_run(str(text))
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = _rgb(color)
    return p


def _caption_p(doc, text) -> None:
    _para(doc, str(text).upper(), bold=True, color=MUTED, size=8.5, space_before=2, space_after=2)


def _divider(doc, eyebrow, title, description, *, key=None, accent=BLUE, bg_default=INK) -> None:
    """Section intro: a full-width colour banner (mirrors the PDF's divider pages),
    themed per-section via section_styles when `key` is given."""
    doc.add_page_break()
    bg = _ss_col(key, "bg", bg_default) if key else None
    acc = _ss_col(key, "accent", accent) if key else accent
    txt = _ss_col(key, "text", WHITE) if key else WHITE
    if bg is not None:
        t = doc.add_table(rows=1, cols=1)
        _no_borders(t)
        cell = t.cell(0, 0)
        _shade(cell, bg)
        p0 = cell.paragraphs[0]
        r0 = p0.add_run(str(eyebrow).upper()); r0.bold = True; r0.font.size = Pt(11); r0.font.color.rgb = _rgb(acc)
        p1 = cell.add_paragraph()
        r1 = p1.add_run(title); r1.bold = True; r1.font.size = Pt(26); r1.font.color.rgb = _rgb(txt)
        p2 = cell.add_paragraph()
        r2 = p2.add_run(description); r2.font.size = Pt(11.5); r2.font.color.rgb = _rgb((210, 212, 218))
        p2.paragraph_format.space_after = Pt(8)
    else:
        _para(doc, str(eyebrow).upper(), bold=True, color=accent, size=11, space_after=2)
        _heading(doc, title, size=26, color=INK, space_before=2, space_after=6)
        _para(doc, description, color=MUTED, size=11.5, space_after=6)


def _pic(doc, png, width_in, *, center=True) -> None:
    if not png:
        return
    try:
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(png), width=Inches(width_in))
    except Exception as e:  # noqa: BLE001
        log.warning("[RENDER_IMG_ERR] docx chart: %s", e)


def _kpi_row(doc, cards) -> None:
    """cards: [(label, value, rgb)] -> one row of solid coloured stat cells."""
    if not cards:
        return
    table = doc.add_table(rows=1, cols=len(cards))
    _no_borders(table)
    for i, (label, value, rgb) in enumerate(cards):
        cell = table.cell(0, i)
        _shade(cell, rgb)
        pv = cell.paragraphs[0]
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = pv.add_run(str(value))
        rv.bold = True
        rv.font.size = Pt(22)
        rv.font.color.rgb = _rgb(WHITE)
        pl = cell.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = pl.add_run(str(label).upper())
        rl.font.size = Pt(8)
        rl.font.color.rgb = _rgb(WHITE)


def _docx_scorecard(doc, score, g, s, p) -> None:
    """Per-building KPI strip: score + Good/Satisfactory/Unsatisfactory/Total."""
    acc = _sc_rgb(score) if score is not None else BLUE
    _kpi_row(doc, [
        (f"Score /100", score if score is not None else "-", acc),
        ("Good", g, GREEN), ("Satisfactory", s, AMBER),
        ("Unsatisfactory", p, RED), ("Total items", g + s + p, INK),
    ])


def _docx_contents(doc, names) -> None:
    """Front-of-report contents list (Word paginates dynamically, so names only)."""
    if not names:
        return
    _para(doc, "IN THIS REPORT", bold=True, color=BLUE, size=11, space_after=2)
    _heading(doc, "Contents", size=26, color=INK, space_before=0, space_after=8)
    for n in names:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(n)
        r.font.size = Pt(13)
        r.font.color.rgb = _rgb(INK)
    doc.add_page_break()


# ---- DOCX section builders ----
def _docx_cover(doc, survey, content) -> None:
    _para(doc, BRAND.upper(), bold=True, color=INK, size=12, space_after=24)
    _heading(doc, "Facility Health", size=34, color=INK, space_before=30, space_after=0)
    p = doc.add_paragraph()
    r1 = p.add_run("Report ")
    r1.bold = True
    r1.font.size = Pt(34)
    r1.font.color.rgb = _rgb(INK)
    r2 = p.add_run(str(datetime.now().year))
    r2.bold = True
    r2.font.size = Pt(34)
    r2.font.color.rgb = _rgb(BLUE)
    _para(doc, survey.get("facility_name") or "Facility", color=MUTED, size=16, space_after=12)
    lines = _facility_lines(survey)
    table = doc.add_table(rows=len(lines), cols=2)
    _no_borders(table)
    for i, (label, val) in enumerate(lines):
        _cell(table.cell(i, 0), label.upper(), bold=True, color=MUTED, size=9, fill=CREAM)
        _cell(table.cell(i, 1), val, color=INK, size=11, fill=CREAM)


def _docx_exec(doc, content, groups) -> None:
    _heading(doc, "Executive Summary", size=20, color=INK, accent=BLUE)
    counts = _rating_counts(groups)
    _pic(doc, charts.rating_donut(counts["good"], counts["satis"], counts["poor"], counts["na"]), 3.0)
    _caption_p(doc, "Rating distribution")
    _para(doc, content.executive_summary, color=INK, size=10.5, space_after=6)
    grade = content.overall_rating
    t = doc.add_table(rows=1, cols=1)
    _no_borders(t)
    _cell(t.cell(0, 0), f"OVERALL RATING:   {str(grade).upper()}", bold=True, color=WHITE,
          size=13, align="center", fill=_grade_rgb(grade))


def _docx_health(doc, health, content) -> None:
    score = (health or {}).get("overall")
    grade = (health or {}).get("grade") or content.overall_rating
    _heading(doc, "Facility Health Score", size=20, color=INK,
             accent=_sc_rgb(score) if score is not None else BLUE)
    _pic(doc, charts.gauge(score, grade), 2.4)
    _caption_p(doc, "Overall health score")
    rows = [(_dl(d["domain"]), d["score"], d.get("graded", 0)) for d in (health or {}).get("domains", [])]
    if rows:
        _para(doc, "Each category scored 0-100 from the surveyor's ratings. Longer green bars are "
                   "healthier; short red bars need attention.", color=MUTED, size=9.5)
        _pic(doc, charts.category_bars(rows), 6.4)
        _caption_p(doc, "Health score by category")


def _docx_findings(doc, title, items, accent) -> None:
    if not items:
        return
    _para(doc, title, bold=True, color=accent, size=10, space_before=2, space_after=1)
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(str(it))
        r.font.size = Pt(9.5)
        r.font.color.rgb = _rgb(INK)


def _docx_category(doc, area, domain, answers, sec, photos_for) -> None:
    score = _group_score(answers)
    facts = [a for a in answers if grade_value(a.get("value")) is None and (a.get("value") or a.get("remark"))]
    graded = [a for a in answers if grade_value(a.get("value")) is not None]
    g, s, p_ = _status_counts(answers)
    title = _dl(domain) + (f"    {score}/100" if score is not None else "")
    _heading(doc, title, size=14, color=INK,
             accent=_sc_rgb(score) if score is not None else MUTED, space_before=8)

    if facts:
        _caption_p(doc, "Key facts")
        t = doc.add_table(rows=len(facts), cols=2)
        _no_borders(t)
        for i, a in enumerate(facts):
            _cell(t.cell(i, 0), a.get("question", ""), bold=True, color=INK, size=9.5, fill=CREAM)
            _cell(t.cell(i, 1), str(a.get("value") or a.get("remark") or ""), color=INK, size=9.5, fill=CREAM)

    if graded:
        _caption_p(doc, "Status summary")
        _para(doc, f"{g} Good      {s} Satisfactory      {p_} Unsatisfactory",
              color=MUTED, size=9.5, space_after=2)
        flagged = [a for a in graded if grade_value(a.get("value")) in (0.0, 0.5)]
        if flagged:
            _para(doc, "Needs attention", bold=True, color=INK, size=9.5, space_after=1)
            t = doc.add_table(rows=len(flagged), cols=2)
            _no_borders(t)
            for i, a in enumerate(flagged):
                _cell(t.cell(i, 0), str(a.get("value")), bold=True, color=WHITE, size=9,
                      align="center", fill=_rating_rgb(a.get("value")))
                _cell(t.cell(i, 1), a.get("question", ""), color=INK, size=9.5)
        if g:
            _para(doc, f"+ {g} item(s) rated Good - no action required.",
                  italic=True, color=MUTED, size=9, space_after=2)

    if sec:
        _docx_findings(doc, "Risks", sec.risks, RED)
        _docx_findings(doc, "Recommendations", sec.recommendations, BLUE)
    _docx_photos(doc, photos_for)


def _docx_corrective(doc, actions) -> None:
    if not actions:
        return
    high = [a for a in actions if a.get("severity") == "high"]
    med = [a for a in actions if a.get("severity") == "medium"]
    _heading(doc, "Priority overview", size=16, color=INK, accent=RED)
    _para(doc, "Every deficiency found during the survey is captured below and prioritised by urgency. "
               "HIGH-priority items are safety- or resident-experience critical and should be closed "
               "immediately; MEDIUM-priority items are planned improvements.", color=INK, size=10)
    _kpi_row(doc, [("Total actions", len(actions), BLUE), ("High priority", len(high), RED),
                   ("Medium priority", len(med), AMBER)])
    _pic(doc, charts.action_priority(len(high), len(med)), 4.2)

    _heading(doc, "Remediation roadmap", size=14, color=INK, accent=BLUE)
    _para(doc, "A recommended schedule for closing each group of issues over the coming weeks.",
          color=MUTED, size=9.5)
    _pic(doc, charts.gantt(_gantt_rows(actions)), 6.4)
    _caption_p(doc, "Timeline by category  ·  weeks from now")
    _para(doc, "Red = immediate (high priority, weeks 0-2).   Amber = planned (medium priority).   "
               "Blue dashed line = recommended re-inspection.", italic=True, color=MUTED, size=9)

    _heading(doc, "Action detail", size=14, color=INK, accent=INK)
    for i, ac in enumerate(actions, 1):
        sev = ac.get("severity", "medium")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        tag = p.add_run(f" {sev.upper()} ")
        tag.bold = True
        tag.font.size = Pt(9)
        tag.font.color.rgb = _rgb(WHITE)
        _run_shade(tag, RED if sev == "high" else AMBER)
        head = p.add_run(f"   {i}.  {_dl(ac.get('domain', ''))}  ·  {ac.get('area', '')}")
        head.bold = True
        head.font.size = Pt(10.5)
        head.font.color.rgb = _rgb(INK)
        _para(doc, f"Finding: {ac.get('finding')} - {ac.get('question')}", color=INK, size=9.5, space_after=1)
        _para(doc, f"Action: {ac.get('action')}", italic=True, color=MUTED, size=9.5, space_after=5)


def _docx_key_recs(doc, content) -> None:
    for i, rec in enumerate(content.key_recommendations, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        num = p.add_run(f"{i}   ")
        num.bold = True
        num.font.size = Pt(12)
        num.font.color.rgb = _rgb(BLUE)
        r = p.add_run(str(rec))
        r.font.size = Pt(11)
        r.font.color.rgb = _rgb(INK)


def _docx_photos(doc, items) -> None:
    """3-per-row photo grid with small captions (mirrors the PDF photo grid)."""
    items = [it for it in (items or []) if it.get("data")]
    if not items:
        return
    _para(doc, "Photos", bold=True, color=INK, size=10, space_before=2, space_after=1)
    cols = 3
    rows = (len(items) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    _no_borders(table)
    for idx, it in enumerate(items):
        cell = table.cell(idx // cols, idx % cols)
        try:
            cell.paragraphs[0].add_run().add_picture(io.BytesIO(it["data"]), width=Inches(2.0))
            cap = cell.add_paragraph()
            rc = cap.add_run(str(it.get("question") or "Photo")[:70])
            rc.font.size = Pt(7)
            rc.font.color.rgb = _rgb(MUTED)
        except Exception as e:  # noqa: BLE001
            log.warning("[RENDER_IMG_ERR] docx photo: %s", e)


def _group(sections, primary):
    from collections import OrderedDict
    out: "OrderedDict[str, list]" = OrderedDict()
    for s in sections:
        ok = s.domain if primary == "domain" else s.area
        ik = s.area if primary == "domain" else s.domain
        out.setdefault(ok, []).append((ik, s))
    return out


def _docx_body(doc, sections, photos, primary, with_photos):
    for outer, inners in _group(sections, primary).items():
        outer_label = _dl(outer) if primary == "domain" else outer
        doc.add_heading(outer_label, level=1)
        for inner, sec in inners:
            inner_label = inner if primary == "domain" else _dl(inner)
            doc.add_heading(inner_label, level=2)
            for label, items in (("Observations", sec.observations), ("Risks", sec.risks),
                                 ("Recommendations", sec.recommendations)):
                if not items:
                    continue
                doc.add_heading(label, level=3)
                for it in items:
                    doc.add_paragraph(it, style="List Bullet")
            if with_photos:
                _docx_photos(doc, photos.get(_pk(sec.area, sec.domain), []))


def render_docx(content: ReportContent, survey: dict,
                photos: Optional[PhotosMap] = None, view: str = "domain",
                health: Optional[dict] = None, actions: Optional[list] = None,
                groups: Optional[list] = None, template: Optional[dict] = None) -> bytes:
    """Word report mirroring the editorial PDF. Applies the same admin template
    (colours, category labels, section text) under the shared render lock."""
    with _render_lock:
        try:
            apply_template(template)
            return _render_docx_body(content, survey, photos or {}, view, health, actions, groups or [])
        finally:
            _reset_template()


def _render_docx_body(content, survey, photos, view, health, actions, groups) -> bytes:
    doc = Document()

    _docx_cover(doc, survey, content)
    doc.add_page_break()

    grp = {(g["area"], g["domain"]): g.get("answers", []) for g in groups}
    secidx = {(s.area, s.domain): s for s in content.sections}
    areas = _ordered_areas({a for (a, _d) in grp.keys()})

    # Contents (names only — Word repaginates on open)
    toc_names = ["Executive Summary", "Facility Health Score"]
    if areas:
        toc_names.append(SECTIONS["buildings"][1])
    if actions:
        toc_names.append(SECTIONS["corrective"][1])
    if content.key_recommendations:
        toc_names.append(SECTIONS["key_recs"][1])
    _docx_contents(doc, toc_names)

    _docx_exec(doc, content, groups)
    _docx_health(doc, health, content)

    if areas:
        # ---- rich per-building path (uses raw answers + AI findings) ----
        _divider(doc, *SECTIONS["buildings"], key="buildings")
        idx = 0
        for area in areas:
            doms = [d for (a, d) in grp.keys() if a == area]
            if area != "Site Profile":
                doms = [d for d in doms if d != "general"]
            if not doms:
                continue
            idx += 1
            is_overview = area == "Site Profile"
            name = "Facility Overview" if is_overview else area
            ss_key = "facility_overview" if is_overview else "building"
            bscore = _group_score([a for d in doms for a in grp.get((area, d), [])])
            _para(doc, f"BUILDING {idx:02d}", bold=True, color=_ss_col(ss_key, "accent", BLUE), size=10,
                  space_before=10, space_after=0)
            _heading(doc, name + (f"    {bscore}/100" if bscore is not None else ""),
                     size=22, color=INK, space_before=0, space_after=4)
            bg_g = bg_s = bg_p = 0
            for d in doms:
                gg, ss, pp = _status_counts(grp.get((area, d), []))
                bg_g += gg; bg_s += ss; bg_p += pp
            if (bg_g + bg_s + bg_p) > 0:
                _docx_scorecard(doc, bscore, bg_g, bg_s, bg_p)
            for d in doms:
                _docx_category(doc, area, d, grp.get((area, d), []), secidx.get((area, d)),
                               photos.get(_pk(area, d), []))
    else:
        # ---- fallback: no per-area answers -> AI narrative sections ----
        _divider(doc, *SECTIONS["buildings"], key="buildings")
        _docx_body(doc, content.sections, photos, "area" if view == "area" else "domain",
                   with_photos=True)

    if actions:
        _divider(doc, *SECTIONS["corrective"], key="corrective")
        _docx_corrective(doc, actions)

    grids = _deploy_rows(survey.get("deployment_plan") or {})
    if grids:
        _heading(doc, "Staff Deployment Plan", size=18, color=INK, accent=BLUE, space_before=12)
        for title, lines in grids:
            _heading(doc, title, size=13, color=INK, accent=MUTED)
            t = doc.add_table(rows=len(lines), cols=2)
            _no_borders(t)
            for i, (role, txt) in enumerate(lines):
                _cell(t.cell(i, 0), role, bold=True, color=INK, size=9.5, fill=CREAM)
                _cell(t.cell(i, 1), txt, color=INK, size=9.5, fill=CREAM)

    if content.key_recommendations:
        _divider(doc, *SECTIONS["key_recs"], accent=LIME_G)
        _docx_key_recs(doc, content)

    _para(doc, BACK_COVER_LINE, italic=True, color=MUTED, size=10, space_before=16, space_after=2)
    _para(doc, f"{BRAND}  ·  AI Facility Intelligence", bold=True, color=BLUE, size=10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
